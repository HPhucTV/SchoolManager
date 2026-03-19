import os
import sys
import subprocess

def install_and_import(package):
    try:
        __import__(package)
    except ImportError:
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])

install_and_import('docx')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml


def add_action(doc, text):
    """Hành động demo — in đậm, nền vàng nhạt."""
    p = doc.add_paragraph()
    run = p.add_run(f'👉 [THAO TÁC] {text}')
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0xBF, 0x36, 0x00)
    run.font.name = 'Arial'
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    # Light yellow background
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="FFF8E1"/>')
    p._p.get_or_add_pPr().append(shading)
    return p


def add_speech(doc, text):
    """Câu nói — font lớn, dễ đọc."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(13)
    run.font.name = 'Arial'
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.35
    return p


def add_transition(doc, text):
    """Câu chuyển đoạn — in nghiêng."""
    p = doc.add_paragraph()
    run = p.add_run(f'⏩ {text}')
    run.italic = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)
    run.font.name = 'Arial'
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    return p


def add_section_title(doc, text, color_hex="1A73E8"):
    """Tiêu đề phần lớn, dễ nhìn."""
    heading = doc.add_heading(text, level=1)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(
            int(color_hex[0:2], 16),
            int(color_hex[2:4], 16),
            int(color_hex[4:6], 16)
        )
        run.font.size = Pt(18)
    return heading


def add_sub_title(doc, text, color_hex="0D47A1"):
    """Tiêu đề phụ."""
    heading = doc.add_heading(text, level=2)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(
            int(color_hex[0:2], 16),
            int(color_hex[2:4], 16),
            int(color_hex[4:6], 16)
        )
        run.font.size = Pt(15)
    return heading


def add_tip(doc, text):
    """Ghi chú nhỏ cho bản thân."""
    p = doc.add_paragraph()
    run = p.add_run(f'💡 {text}')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.italic = True
    run.font.name = 'Arial'
    p.paragraph_format.space_after = Pt(2)
    return p


def add_separator(doc):
    p = doc.add_paragraph()
    run = p.add_run('─' * 60)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)


def create_script():
    doc = Document()

    # Page setup
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Default style
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(13)
    style.paragraph_format.line_spacing = 1.35
    style.paragraph_format.space_after = Pt(4)

    # ═══════════════════════════════════════════════════════════
    # TRANG BÌA
    # ═══════════════════════════════════════════════════════════
    for _ in range(5):
        doc.add_paragraph()

    title = doc.add_heading('BÀI THUYẾT TRÌNH DEMO', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_heading('DỰ ÁN SCHOOLMANAGER', 0)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)

    doc.add_paragraph()
    p = doc.add_paragraph('Kịch bản thuyết trình — Đọc theo thứ tự từ trên xuống')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.italic = True
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Quy ước đọc:')
    run.bold = True
    run.font.size = Pt(12)

    rules = [
        '📖 Chữ thường = CÂU ĐỌC (đọc to cho giám khảo nghe)',
        '👉 [THAO TÁC] Nền vàng = HÀNH ĐỘNG trên màn hình (không đọc, chỉ làm)',
        '💡 Chữ nhỏ xám = GHI CHÚ cho bản thân (không đọc)',
        '⏩ Chữ nghiêng xanh = CÂU CHUYỂN ĐOẠN',
    ]
    for rule in rules:
        p = doc.add_paragraph(rule)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.size = Pt(11)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # PHẦN MỞ ĐẦU
    # ═══════════════════════════════════════════════════════════
    add_section_title(doc, '🎤 PHẦN MỞ ĐẦU — GIỚI THIỆU DỰ ÁN')

    add_action(doc, 'Mở trang Landing Page: https://schoolmanager.id.vn')

    add_speech(doc,
        'Xin chào thầy cô và ban giám khảo. '
        'Hôm nay em xin trình bày dự án SchoolManager — '
        'một nền tảng quản lý trường học thông minh tích hợp trí tuệ nhân tạo.')

    add_speech(doc,
        'Dự án được xây dựng với mục tiêu chuyển đổi số cho giáo dục, '
        'giúp kết nối ba đối tượng: Quản trị viên, Giáo viên và Học sinh '
        'trên cùng một hệ thống.')

    add_speech(doc,
        'SchoolManager không chỉ quản lý truyền thống như điểm số hay thời khóa biểu, '
        'mà còn tích hợp bốn yếu tố nổi bật:')

    add_speech(doc,
        'Thứ nhất là AI — trí tuệ nhân tạo xuyên suốt hệ thống, '
        'từ chatbot hỗ trợ, tạo đề thi tự động, đến gia sư ảo cá nhân hóa.')

    add_speech(doc,
        'Thứ hai là Gamification — trò chơi hóa việc học, '
        'giúp học sinh có động lực với hệ thống điểm, huy hiệu và bảng xếp hạng.')

    add_speech(doc,
        'Thứ ba là Wellness — chăm sóc sức khỏe tinh thần, '
        'theo dõi cảm xúc học sinh hàng ngày và có cơ chế cảnh báo SOS khi cần.')

    add_speech(doc,
        'Thứ tư là Smart Search — hệ thống tìm kiếm thông minh '
        'sử dụng thuật toán BM25, giúp tìm kiếm nhanh chóng trong toàn bộ hệ thống.')

    add_speech(doc,
        'Về công nghệ, dự án sử dụng Next.js và React cho frontend, '
        'FastAPI với Python cho backend, '
        'và được triển khai bằng Docker trên server thực tế.')

    add_transition(doc, 'Bây giờ em xin demo trực tiếp, bắt đầu từ giao diện Quản trị viên.')

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # PHẦN 1: ADMIN
    # ═══════════════════════════════════════════════════════════
    add_section_title(doc, '🏛️ PHẦN 1 — QUẢN TRỊ VIÊN (ADMIN)', 'E65100')

    add_action(doc, 'Đăng nhập bằng tài khoản admin@happyschools.vn / test123')

    add_speech(doc,
        'Đầu tiên em sẽ demo giao diện Quản trị viên. '
        'Admin là người quản lý toàn bộ hệ thống, '
        'có toàn quyền thêm, sửa, xóa dữ liệu.')

    add_separator(doc)

    # Admin Dashboard
    add_sub_title(doc, '1.1. Dashboard Quản trị')
    add_action(doc, 'Chỉ vào trang Dashboard Admin')

    add_speech(doc,
        'Đây là Dashboard tổng quan. '
        'Admin có thể thấy ngay các con số quan trọng: '
        'tổng số học sinh, giáo viên, lớp học trong hệ thống. '
        'Giao diện được thiết kế trực quan để nắm bắt tình hình nhanh nhất.')

    add_separator(doc)

    # Quản lý học sinh
    add_sub_title(doc, '1.2. Quản lý Học sinh')
    add_action(doc, 'Click vào menu "Quản lý Học sinh"')

    add_speech(doc,
        'Mục Quản lý Học sinh cho phép Admin xem toàn bộ danh sách học sinh, '
        'tìm kiếm, lọc theo lớp. '
        'Admin có thể thêm học sinh mới, chỉnh sửa thông tin, '
        'hoặc xóa tài khoản khi cần.')

    add_action(doc, 'Chỉ vào nút Thêm học sinh, các nút Sửa / Xóa')

    add_speech(doc,
        'Ngoài ra, hệ thống còn hỗ trợ import danh sách học sinh từ file Excel, '
        'giúp tiết kiệm thời gian khi nhập số lượng lớn.')

    add_separator(doc)

    # Quản lý giáo viên
    add_sub_title(doc, '1.3. Quản lý Giáo viên')
    add_action(doc, 'Click vào menu "Quản lý Giáo viên"')

    add_speech(doc,
        'Tương tự, trang Quản lý Giáo viên cho phép Admin '
        'thêm mới, chỉnh sửa và xóa tài khoản giáo viên. '
        'Mỗi giáo viên sau khi được tạo tài khoản '
        'sẽ có thể đăng nhập và quản lý lớp của mình.')

    add_separator(doc)

    # Quản lý lớp học
    add_sub_title(doc, '1.4. Quản lý Lớp học')
    add_action(doc, 'Click vào menu "Quản lý Lớp học"')

    add_speech(doc,
        'Admin quản lý danh sách lớp học, '
        'tạo lớp mới, gán giáo viên chủ nhiệm cho từng lớp, '
        'và phân bổ học sinh vào lớp.')

    add_tip(doc, 'Nếu có thời gian, bấm tạo 1 lớp mới để demo nhanh.')

    add_transition(doc, 'Đó là toàn bộ chức năng của Admin. Tiếp theo, em sẽ chuyển sang giao diện Giáo viên, nơi có nhiều tính năng phong phú hơn.')

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # PHẦN 2: GIÁO VIÊN
    # ═══════════════════════════════════════════════════════════
    add_section_title(doc, '👩‍🏫 PHẦN 2 — GIÁO VIÊN (TEACHER)', '1565C0')

    add_action(doc, 'Đăng xuất Admin. Đăng nhập bằng gv.thao@happyschools.vn / test123')

    add_speech(doc,
        'Bây giờ em sẽ demo giao diện dành cho Giáo viên. '
        'Đây là vai trò có nhiều chức năng nhất trong hệ thống, '
        'với tổng cộng 13 trang chức năng.')

    add_separator(doc)

    # Teacher Dashboard
    add_sub_title(doc, '2.1. Dashboard Giáo viên')
    add_action(doc, 'Chỉ vào trang Dashboard')

    add_speech(doc,
        'Dashboard giáo viên hiển thị tổng quan nhanh: '
        'số lớp đang quản lý, số học sinh, '
        'bài tập và kiểm tra gần đây, '
        'cùng các hoạt động mới nhất.')

    add_separator(doc)

    # Quản lý lớp học
    add_sub_title(doc, '2.2. Quản lý Lớp học')
    add_action(doc, 'Click menu "Lớp học"')

    add_speech(doc,
        'Giáo viên xem danh sách lớp mình phụ trách. '
        'Có thể click vào từng lớp để xem chi tiết: '
        'danh sách học sinh, mời thêm học sinh vào lớp, '
        'và xem các hoạt động của lớp.')

    add_action(doc, 'Click vào 1 lớp để xem chi tiết')

    add_speech(doc,
        'Trong chi tiết lớp, giáo viên thấy danh sách đầy đủ '
        'và có thể gửi lời mời tham gia lớp cho học sinh mới.')

    add_separator(doc)

    # Thời khóa biểu
    add_sub_title(doc, '2.3. Thời khóa biểu')
    add_action(doc, 'Click menu "Thời khóa biểu"')

    add_speech(doc,
        'Trang Thời khóa biểu hiển thị lịch dạy của giáo viên '
        'theo dạng lưới tuần, rất trực quan. '
        'Giáo viên có thể thêm, sửa, xóa các tiết học.')

    add_separator(doc)

    # KIỂM TRA — Điểm nhấn AI
    add_sub_title(doc, '2.4. Tạo bài kiểm tra — Tính năng AI nổi bật ⭐')
    add_action(doc, 'Click menu "Kiểm tra" → Tạo bài kiểm tra mới')

    add_speech(doc,
        'Đây là một trong những tính năng quan trọng nhất. '
        'Giáo viên có BA cách tạo đề kiểm tra:')

    add_speech(doc,
        'Cách 1: Tạo thủ công — giáo viên tự nhập từng câu hỏi và đáp án.')

    add_speech(doc,
        'Cách 2: Tạo bằng AI — chỉ cần nhập chủ đề và môn học, '
        'AI sẽ tự động tạo toàn bộ câu hỏi trắc nghiệm. '
        'Giáo viên có thể chọn mức độ khó: dễ, trung bình, hoặc khó.')

    add_action(doc, 'Demo tạo đề bằng AI: nhập chủ đề, chọn môn, bấm Tạo')
    add_tip(doc, 'Ví dụ: Chủ đề "Phương trình bậc 2", Môn "Toán", Độ khó "Trung bình", 5 câu')

    add_speech(doc,
        'Như các thầy cô thấy, chỉ trong vài giây, '
        'AI đã tự động tạo ra bộ đề hoàn chỉnh. '
        'Giáo viên vẫn có thể chỉnh sửa lại trước khi công bố.')

    add_speech(doc,
        'Cách 3: Import từ file Word — giáo viên upload file .docx '
        'có sẵn đề thi, hệ thống sẽ tự động nhận diện và phân tích '
        'thành các câu hỏi.')

    add_speech(doc,
        'Ngoài ra, giáo viên có thể đặt deadline cho bài kiểm tra. '
        'Sau deadline, hệ thống sẽ tự động khóa, học sinh không thể nộp thêm.')

    add_separator(doc)

    # Bài tập
    add_sub_title(doc, '2.5. Quản lý Bài tập')
    add_action(doc, 'Click menu "Bài tập"')

    add_speech(doc,
        'Giáo viên tạo bài tập với mô tả chi tiết, '
        'điểm tối đa và deadline. '
        'Học sinh sẽ nộp bài trực tuyến.')

    add_speech(doc,
        'Khi chấm bài, giáo viên có hai lựa chọn: '
        'chấm thủ công hoặc nhờ AI chấm tự động. '
        'AI sẽ phân tích nội dung bài làm và đưa ra nhận xét chi tiết, '
        'giúp giáo viên tiết kiệm rất nhiều thời gian.')

    add_action(doc, 'Chỉ vào nút "Chấm bằng AI" nếu có bài nộp')

    add_separator(doc)

    # Quiz Battle
    add_sub_title(doc, '2.6. Quiz Battle — Đấu trường kiến thức ⚔️')
    add_action(doc, 'Click menu "Quiz Battle"')

    add_speech(doc,
        'Quiz Battle là chế độ thi đấu thời gian thực '
        'giữa các học sinh. '
        'Giáo viên tạo phòng thi với mã phòng riêng, '
        'học sinh nhập mã để tham gia.')

    add_speech(doc,
        'Khi trận đấu bắt đầu, tất cả học sinh cùng trả lời câu hỏi. '
        'Điểm số được cập nhật realtime trên bảng xếp hạng, '
        'tạo sự cạnh tranh rất hứng thú.')

    add_separator(doc)

    # Hoạt động
    add_sub_title(doc, '2.7. Quản lý Hoạt động')
    add_action(doc, 'Click menu "Hoạt động"')

    add_speech(doc,
        'Giáo viên tạo và quản lý các hoạt động ngoại khóa, '
        'sự kiện của lớp. Học sinh sẽ thấy được các hoạt động này '
        'trên giao diện của mình.')

    add_separator(doc)

    # Sức khỏe tinh thần
    add_sub_title(doc, '2.8. Sức khỏe tinh thần — Wellness 💚')
    add_action(doc, 'Click menu "Sức khỏe"')

    add_speech(doc,
        'Đây là tính năng rất nhân văn của dự án. '
        'Giáo viên có thể theo dõi sức khỏe tinh thần '
        'của toàn bộ học sinh trong lớp.')

    add_speech(doc,
        'Trang này hiển thị biểu đồ phân bố cảm xúc, '
        'nhận diện xu hướng tiêu cực, '
        'và quan trọng nhất là nhận được cảnh báo SOS '
        'khi có học sinh cần hỗ trợ khẩn cấp.')

    add_speech(doc,
        'Tính năng SOS cho phép học sinh gửi tín hiệu cầu cứu, '
        'kể cả ẩn danh, khi gặp vấn đề về tâm lý. '
        'Giáo viên sẽ được thông báo ngay lập tức để can thiệp kịp thời.')

    add_separator(doc)

    # Phân tích & Thống kê
    add_sub_title(doc, '2.9. Phân tích & Thống kê 📊')
    add_action(doc, 'Click menu "Phân tích" hoặc "Thống kê"')

    add_speech(doc,
        'Hệ thống cung cấp Dashboard phân tích toàn diện. '
        'Giáo viên xem được phân bố điểm số, '
        'tỷ lệ hoàn thành bài tập, '
        'xu hướng phát triển của từng học sinh.')

    add_speech(doc,
        'Đặc biệt có tính năng Cảnh báo sớm, '
        'tự động phát hiện học sinh có điểm giảm sút liên tục, '
        'vắng mặt nhiều, hoặc có dấu hiệu bất ổn, '
        'để giáo viên hỗ trợ kịp thời.')

    add_separator(doc)

    # Thi đua
    add_sub_title(doc, '2.10. Thi đua & Gamification 🏆')
    add_action(doc, 'Click menu "Thi đua"')

    add_speech(doc,
        'Giáo viên xem bảng xếp hạng lớp, '
        'quản lý hệ thống huy hiệu và cửa hàng phần thưởng. '
        'Đây là nơi giáo viên thiết lập cơ chế Gamification cho lớp mình.')

    add_separator(doc)

    # Chatbot giáo viên
    add_sub_title(doc, '2.11. Trợ lý AI Chatbot 🤖')
    add_action(doc, 'Click vào biểu tượng Chatbot ở góc phải dưới')

    add_speech(doc,
        'Giáo viên có trợ lý AI riêng. '
        'Chatbot có thể giúp soạn báo cáo lớp, '
        'lên kế hoạch bài giảng, '
        'phân tích kết quả học tập, '
        'hoặc gợi ý phương pháp giảng dạy.')

    add_action(doc, 'Gõ thử 1 câu: "Giúp tôi tạo báo cáo tổng kết lớp 10A"')
    add_tip(doc, 'Chờ AI trả lời, chỉ cho giám khảo xem kết quả.')

    add_separator(doc)

    # Smart Search
    add_sub_title(doc, '2.12. Tìm kiếm thông minh — Smart Search 🔍')
    add_action(doc, 'Bấm Ctrl+K hoặc click thanh tìm kiếm trên Header')

    add_speech(doc,
        'Hệ thống có công cụ tìm kiếm toàn cục, '
        'mở bằng phím tắt Ctrl+K. '
        'Tìm kiếm được xuyên suốt: học sinh, lớp, bài tập, kiểm tra, '
        'hoạt động, thông báo.')

    add_speech(doc,
        'Thuật toán sử dụng BM25 kết hợp AI Ranking, '
        'hỗ trợ tìm kiếm tiếng Việt không dấu. '
        'Ví dụ gõ "bai tap" vẫn tìm được "bài tập".')

    add_action(doc, 'Gõ thử 1 từ khóa, chỉ kết quả cho giám khảo')

    add_transition(doc, 'Vậy là em đã demo xong giao diện Giáo viên. Tiếp theo là giao diện Học sinh — nơi các em học sinh sẽ tương tác hàng ngày.')

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # PHẦN 3: HỌC SINH
    # ═══════════════════════════════════════════════════════════
    add_section_title(doc, '🎓 PHẦN 3 — HỌC SINH (STUDENT)', '2E7D32')

    add_action(doc, 'Đăng xuất Giáo viên. Đăng nhập bằng hs.an0@happyschools.vn / test123')

    add_speech(doc,
        'Cuối cùng là giao diện Học sinh — '
        'nơi các em sẽ tương tác hàng ngày. '
        'Giao diện được thiết kế trẻ trung, thân thiện '
        'và có rất nhiều yếu tố tạo động lực học tập.')

    add_separator(doc)

    # Student Dashboard
    add_sub_title(doc, '3.1. Dashboard Học sinh')
    add_action(doc, 'Chỉ vào trang Dashboard')

    add_speech(doc,
        'Dashboard học sinh hiển thị tổng quan cá nhân: '
        'level hiện tại, điểm XP, chuỗi ngày điểm danh, '
        'huy hiệu đã đạt được, bài tập sắp deadline, '
        'và các hoạt động gần đây.')

    add_speech(doc,
        'Ngay từ Dashboard, học sinh đã thấy được '
        'tính năng Gamification với level, XP và huy hiệu — '
        'tạo cảm giác như đang chơi game RPG.')

    add_separator(doc)

    # Check-in
    add_sub_title(doc, '3.2. Daily Check-in ✨')
    add_action(doc, 'Bấm nút Check-in trên Dashboard (nếu chưa check-in hôm nay)')

    add_speech(doc,
        'Mỗi ngày, học sinh check-in để nhận điểm XP. '
        'Hệ thống theo dõi chuỗi ngày liên tục — '
        'giống streak trong Duolingo, '
        'khuyến khích học sinh đăng nhập đều đặn.')

    add_separator(doc)

    # Lớp học
    add_sub_title(doc, '3.3. Lớp học')
    add_action(doc, 'Click menu "Lớp học"')

    add_speech(doc,
        'Học sinh xem thông tin lớp mình tham gia, '
        'danh sách bạn cùng lớp, '
        'và các bài tập, kiểm tra của lớp.')

    add_separator(doc)

    # Làm quiz
    add_sub_title(doc, '3.4. Làm bài kiểm tra')
    add_action(doc, 'Click vào 1 bài kiểm tra đang mở')

    add_speech(doc,
        'Học sinh vào làm bài kiểm tra trắc nghiệm '
        'với giao diện rõ ràng, dễ sử dụng. '
        'Sau khi nộp bài, điểm số hiển thị ngay lập tức, '
        'kèm theo so sánh với đáp án đúng để học sinh tự rút kinh nghiệm.')

    add_action(doc, 'Nếu có bài quiz, demo làm nhanh 1-2 câu rồi nộp')

    add_separator(doc)

    # Quiz Battle
    add_sub_title(doc, '3.5. Quiz Battle — Đấu trường PvP ⚔️')
    add_action(doc, 'Click menu "Quiz Battle"')

    add_speech(doc,
        'Từ phía học sinh, các em nhập mã phòng do giáo viên tạo '
        'để tham gia đấu trường. '
        'Trải nghiệm thi đấu thời gian thực với bạn bè '
        'tạo sự hứng thú rất lớn.')

    add_separator(doc)

    # Bài tập
    add_sub_title(doc, '3.6. Bài tập')
    add_action(doc, 'Click menu "Bài tập"')

    add_speech(doc,
        'Học sinh xem danh sách bài tập được giáo viên giao. '
        'Với mỗi bài tập có hiển thị deadline, trạng thái, '
        'và học sinh nộp bài trực tuyến ngay trên web.')

    add_separator(doc)

    # AI Tutor
    add_sub_title(doc, '3.7. AI Tutor — Gia sư ảo cá nhân hóa 🤖')
    add_action(doc, 'Click menu "AI Tutor"')

    add_speech(doc,
        'Đây là tính năng đặc biệt của dự án. '
        'AI Tutor đóng vai trò gia sư riêng cho từng học sinh, '
        'tự động phân tích kết quả học tập để đưa ra lộ trình cá nhân.')

    add_speech(doc,
        'AI sẽ phân tích điểm mạnh, điểm yếu của học sinh theo từng môn và chủ đề. '
        'Ví dụ, nếu em yếu Toán ở phần phương trình, '
        'AI sẽ đề xuất ôn tập chủ đề đó ưu tiên trước.')

    add_speech(doc,
        'Hệ thống chia thành 3 giai đoạn: '
        'Nâng cao cho môn giỏi, '
        'Trung bình cho môn cần ôn tập, '
        'và Cần cải thiện cho môn cần tập trung hơn. '
        'Kèm theo đó là lời khuyên học tập phù hợp.')

    add_separator(doc)

    # Mood Journal
    add_sub_title(doc, '3.8. Nhật ký cảm xúc — Mood Journal 💚')
    add_action(doc, 'Click menu "Nhật ký cảm xúc"')

    add_speech(doc,
        'Mỗi ngày, học sinh ghi lại cảm xúc thông qua emoji. '
        'Ví dụ: vui, bình thường, buồn, lo lắng. '
        'Có thể thêm ghi chú về lý do.')

    add_speech(doc,
        'Dữ liệu này được tổng hợp để giáo viên theo dõi. '
        'Và nếu học sinh gặp khó khăn tâm lý, '
        'có thể gửi SOS — kể cả ẩn danh — '
        'để giáo viên kịp thời hỗ trợ.')

    add_action(doc, 'Demo chọn 1 emoji và ghi cảm xúc')
    add_tip(doc, 'Nếu có thời gian, bấm gửi SOS ẩn danh để demo.')

    add_separator(doc)

    # Thành tích
    add_sub_title(doc, '3.9. Thành tích & Huy hiệu 🏅')
    add_action(doc, 'Click menu "Thành tích"')

    add_speech(doc,
        'Trang Thành tích hiển thị bộ sưu tập huy hiệu, '
        'level hiện tại, XP đã tích lũy. '
        'Hệ thống có nhiều loại huy hiệu — '
        'ví dụ: streak 7 ngày, hoàn thành 10 quiz, top 1 bảng xếp hạng.')

    add_speech(doc,
        'Ngoài ra có Bảng xếp hạng để so sánh thành tích '
        'với các bạn trong lớp, '
        'cùng với Cửa hàng phần thưởng để đổi XP lấy quà.')

    add_separator(doc)

    # Mini-games
    add_sub_title(doc, '3.10. Mini-games — Giải trí & Học tập 🎮')
    add_action(doc, 'Click menu "Giải trí"')

    add_speech(doc,
        'SchoolManager có hai mini-game tích hợp: '
        'Riddles — câu đố vui rèn tư duy, '
        'và Word Chain — nối từ tiếng Việt giúp mở rộng vốn từ.')

    add_speech(doc,
        'Khi đạt điểm cao trong game, học sinh cũng nhận được XP, '
        'tạo thêm động lực.')

    add_action(doc, 'Mở 1 game nhanh để demo nếu còn thời gian')

    add_separator(doc)

    # Chatbot học sinh
    add_sub_title(doc, '3.11. AI Chatbot cho Học sinh 🤖')
    add_action(doc, 'Click biểu tượng Chatbot góc phải dưới')

    add_speech(doc,
        'Học sinh có chatbot riêng với khả năng: '
        'giải đáp thắc mắc bài vở, '
        'tư vấn phương pháp học, '
        'và đặc biệt là hỗ trợ tâm lý — '
        'chatbot được huấn luyện để tư vấn sức khỏe tinh thần cho học sinh.')

    add_action(doc, 'Gõ thử: "Tôi cảm thấy lo lắng về bài kiểm tra sắp tới"')
    add_tip(doc, 'Chờ AI trả lời, chỉ cho giám khảo thấy AI phản hồi nhạy cảm, phù hợp.')

    add_separator(doc)

    # Thông báo + Thời khóa biểu + Cài đặt
    add_sub_title(doc, '3.12. Thông báo, Thời khóa biểu & Cài đặt')

    add_speech(doc,
        'Ngoài ra, học sinh còn có: '
        'Thông báo — nhận thông tin về deadline, kết quả kiểm tra, lời mời lớp; '
        'Thời khóa biểu — xem lịch học theo tuần; '
        'và Cài đặt — cập nhật thông tin cá nhân, đổi mật khẩu, tải avatar.')

    add_transition(doc, 'Vậy là em đã demo xong toàn bộ ba giao diện. Em xin chuyển sang phần tổng kết.')

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # PHẦN KẾT
    # ═══════════════════════════════════════════════════════════
    add_section_title(doc, '🎯 PHẦN KẾT — TỔNG KẾT', '7B1FA2')

    add_speech(doc,
        'Tổng kết lại, SchoolManager là nền tảng quản lý trường học toàn diện, '
        'với các điểm nổi bật:')

    add_speech(doc,
        'Một là, trí tuệ nhân tạo xuyên suốt — '
        'từ chatbot, tạo đề thi, chấm bài, '
        'đến gia sư ảo cá nhân hóa cho từng học sinh.')

    add_speech(doc,
        'Hai là, Gamification — biến việc học thành trải nghiệm game RPG, '
        'tạo động lực tự giác cho học sinh mỗi ngày.')

    add_speech(doc,
        'Ba là, chăm sóc sức khỏe tinh thần — '
        'theo dõi cảm xúc, cảnh báo SOS — '
        'đáp ứng nhu cầu cấp thiết trong giáo dục hiện đại.')

    add_speech(doc,
        'Bốn là, kiến trúc chuyên nghiệp — '
        'sử dụng Next.js, FastAPI và Docker, '
        'đạt chuẩn công nghiệp, sẵn sàng triển khai thực tế.')

    add_speech(doc,
        'Dự án có hơn 22 module backend, '
        '29 trang giao diện, '
        'và đã được triển khai hoạt động tại schoolmanager.id.vn.')

    doc.add_paragraph()

    add_speech(doc,
        'Em xin cảm ơn thầy cô và ban giám khảo đã lắng nghe. '
        'Em rất sẵn lòng giải đáp mọi câu hỏi. '
        'Xin cảm ơn ạ!')

    add_tip(doc, 'Cúi đầu cảm ơn. Sẵn sàng trả lời câu hỏi.')

    # ── Save ──
    output_path = os.path.join(
        'c:\\Users\\PC\\Documents\\SchoolManager',
        'BaiThuyetTrinh_Demo_SchoolManager.docx'
    )
    doc.save(output_path)
    print(f"\n✅ Thành công! File Word đã được tạo tại:\n   {output_path}\n")


if __name__ == "__main__":
    create_script()
